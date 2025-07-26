"""
MCP 服务器 DOCX/DOC 文档解析器

该模块提供 Word 文档解析功能，并具有自动将 .doc 转换为 .docx 的能力。
"""

import os
import logging
from typing import Dict, Any, List, Optional

try:
    import docx2txt
except ImportError:
    docx2txt = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

from ..types import ParseResult, FileType, ParserStatus, ConversionMethod
from ..exceptions import ParsingError, EmptyDocumentError, ConversionError
from .base import StructuredParser
from .converters import auto_convert_doc_to_docx
from ..utils import cleanup_temp_path

logger = logging.getLogger(__name__)


class DocxParser(StructuredParser):
    """
    Word 文档解析器 (.docx 和 .doc 文件)。
    
    此解析器处理现代 .docx 文件和旧版 .doc 文件，
    具有自动转换功能。
    """
    
    def __init__(self):
        """初始化 DOCX 解析器。"""
        super().__init__(FileType.DOCX)
        self.supported_extensions = ['.docx', '.doc']
    
    def supports_file(self, file_path: str) -> bool:
        """
        检查此解析器是否支持给定文件。
        
        参数:
            file_path: 要检查的文件路径
            
        返回:
            如果文件是 Word 文档则返回 True，否则返回 False
        """
        return os.path.splitext(file_path)[1].lower() in self.supported_extensions
    
    def parse(self, file_path: str) -> ParseResult:
        """
        解析 Word 文档并提取其内容。
        
        参数:
            file_path: Word 文档路径
            
        返回:
            包含提取内容和元数据的 ParseResult
            
        引发:
            ParsingError: 如果文档解析失败
        """
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_ext == '.doc':
                return self._parse_doc_file(file_path)
            else:
                return self._parse_docx_file(file_path)
                
        except Exception as e:
            logger.error(f"Word 文档解析失败 {file_path}: {str(e)}")
            return self.create_error_result(
                file_path=file_path,
                error=f"Word 文档解析失败: {str(e)}"
            )
    
    def _parse_doc_file(self, file_path: str) -> ParseResult:
        """
        解析旧版 .doc 文件并自动转换。
        
        参数:
            file_path: .doc 文件路径
            
        返回:
            包含提取内容的 ParseResult
        """
        logger.info(f"尝试解析 .doc 文件: {file_path}")
        
        # 首先尝试自动转换
        conversion_result = auto_convert_doc_to_docx(file_path)
        
        if conversion_result.success:
            logger.info(f"使用 {conversion_result.method.value} 成功转换 .doc")
            
            try:
                # 解析转换后的 .docx 文件
                if DocxDocument is not None:
                    result = self._parse_docx_with_python_docx(conversion_result.converted_path)
                else:
                    result = self._parse_docx_with_docx2txt(conversion_result.converted_path)
                
                # 更新元数据以反映原始文件和转换
                if result.success:
                    result.file_path = file_path  # 更新为原始文件路径
                    result.metadata.update({
                        "conversion_method": conversion_result.method.value,
                        "original_format": "doc",
                        "converted_from": "doc",
                        "note": "已成功从 .doc 自动转换为 .docx 格式"
                    })
                
                # 清理临时文件
                if conversion_result.temp_dir:
                    cleanup_temp_path(conversion_result.temp_dir)
                
                return result
                
            except Exception as e:
                # 出错时清理
                if conversion_result.temp_dir:
                    cleanup_temp_path(conversion_result.temp_dir)
                raise e
        else:
            # 转换失败，尝试备用方法
            return self._parse_doc_fallback(file_path, conversion_result)
    
    def _parse_docx_file(self, file_path: str) -> ParseResult:
        """
        使用可用库解析 .docx 文件。
        
        参数:
            file_path: .docx 文件路径
            
        返回:
            包含提取内容的 ParseResult
        """
        # 首先尝试 docx2txt (文本提取更可靠)
        if docx2txt is not None:
            try:
                return self._parse_docx_with_docx2txt(file_path)
            except Exception as e:
                logger.debug(f"docx2txt 解析失败: {e}")
        
        # 备用 python-docx
        if DocxDocument is not None:
            try:
                return self._parse_docx_with_python_docx(file_path)
            except Exception as e:
                logger.debug(f"python-docx 解析失败: {e}")
        
        # 没有可用库
        return self.create_error_result(
            file_path=file_path,
            error="没有可用的 Word 文档解析库。请安装 docx2txt 或 python-docx。",
            status=ParserStatus.ERROR
        )
    
    def _parse_docx_with_docx2txt(self, file_path: str) -> ParseResult:
        """
        使用 docx2txt 库解析 .docx 文件。
        
        参数:
            file_path: .docx 文件路径
            
        返回:
            包含提取内容的 ParseResult
        """
        full_text = docx2txt.process(file_path)
        
        if not full_text or not full_text.strip():
            raise EmptyDocumentError(file_path, "docx2txt")
        
        # 从文本创建结构化内容
        structured_content = self._create_structured_content_from_text(full_text)
        
        metadata = {
            "parsing_method": "docx2txt",
            "file_type": "docx",
            "total_paragraphs": len(structured_content.get("paragraphs", [])),
            "file_size": os.path.getsize(file_path)
        }
        
        return self.create_success_result(
            file_path=file_path,
            content=full_text.strip(),
            metadata=metadata,
            parsing_method="docx2txt"
        )
    
    def _parse_docx_with_python_docx(self, file_path: str) -> ParseResult:
        """
        使用 python-docx 库解析 .docx 文件。
        
        参数:
            file_path: .docx 文件路径
            
        返回:
            包含提取内容的 ParseResult
        """
        doc = DocxDocument(file_path)
        paragraphs = []
        full_text = ""
        
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                paragraph_info = {
                    "paragraph_number": i + 1,
                    "content": paragraph.text.strip(),
                    "style": paragraph.style.name if paragraph.style else "Normal"
                }
                paragraphs.append(paragraph_info)
                full_text += paragraph.text.strip() + "\n\n"
        
        if not full_text.strip():
            raise EmptyDocumentError(file_path, "python-docx")
        
        structured_content = {
            "paragraphs": paragraphs,
            "total_paragraphs": len(paragraphs),
            "full_text": full_text.strip()
        }
        
        metadata = {
            "parsing_method": "python-docx",
            "file_type": "docx",
            "total_paragraphs": len(paragraphs),
            "paragraphs": paragraphs,
            "file_size": os.path.getsize(file_path)
        }
        
        return self.create_success_result(
            file_path=file_path,
            content=full_text.strip(),
            metadata=metadata,
            parsing_method="python-docx"
        )
    
    def _parse_doc_fallback(self, file_path: str, conversion_result) -> ParseResult:
        """
        当转换失败时解析 .doc 文件的备用方法。
        
        参数:
            file_path: .doc 文件路径
            conversion_result: 转换尝试的结果
            
        返回:
            包含错误信息和建议的 ParseResult
        """
        # 直接在 .doc 文件上尝试 docx2txt (有时有效)
        if docx2txt is not None:
            try:
                full_text = docx2txt.process(file_path)
                if full_text and full_text.strip():
                    logger.info("使用 docx2txt 直接成功解析 .doc 文件")
                    
                    structured_content = self._create_structured_content_from_text(full_text)
                    
                    metadata = {
                        "parsing_method": "docx2txt_direct",
                        "file_type": "doc",
                        "total_paragraphs": len(structured_content.get("paragraphs", [])),
                        "file_size": os.path.getsize(file_path),
                        "note": "直接解析 .doc 文件而无需转换"
                    }
                    
                    return self.create_success_result(
                        file_path=file_path,
                        content=full_text.strip(),
                        metadata=metadata,
                        parsing_method="docx2txt_direct"
                    )
            except Exception as e:
                logger.debug(f"使用 docx2txt 直接解析 .doc 失败: {e}")
        
        # 所有方法都失败
        error_msg = "无法解析 .doc 文件"
        suggestions = [
            "使用 Microsoft Word 将 .doc 文件手动转换为 .docx 格式",
            "使用 LibreOffice Writer 打开并导出为 .docx 或 .pdf",
            "尝试在线 DOC 到 DOCX 转换器",
            "确保文件未损坏"
        ]
        
        if conversion_result.error:
            error_msg += f": {conversion_result.error}"
        
        if conversion_result.tried_methods:
            suggestions.append(f"已尝试的转换方法: {', '.join(conversion_result.tried_methods)}")
        
        return self.create_error_result(
            file_path=file_path,
            error=error_msg,
            status=ParserStatus.CONVERSION_FAILED,
            metadata={
                "conversion_error": conversion_result.error,
                "tried_methods": conversion_result.tried_methods,
                "suggestions": suggestions
            }
        )
    
    def _create_structured_content_from_text(self, text: str) -> Dict[str, Any]:
        """
        从纯文本创建结构化内容。
        
        参数:
            text: 纯文本内容
            
        返回:
            包含结构化内容的字典
        """
        # 将文本分割成段落
        lines = text.split('\n')
        paragraphs = []
        paragraph_num = 1
        
        for line in lines:
            line = line.strip()
            if line:  # 非空行
                paragraph_info = {
                    "paragraph_number": paragraph_num,
                    "content": line
                }
                paragraphs.append(paragraph_info)
                paragraph_num += 1
        
        return {
            "paragraphs": paragraphs,
            "total_paragraphs": len(paragraphs),
            "full_text": text
        }
    
    def extract_structured_content(self, file_path: str) -> Dict[str, Any]:
        """
        从 Word 文档中提取结构化内容。
        
        参数:
            file_path: Word 文档路径
            
        返回:
            包含结构化内容的字典
            
        引发:
            ParsingError: 如果提取失败
        """
        parse_result = self.parse(file_path)
        
        if not parse_result.success:
            raise ParsingError(
                message=parse_result.error or "提取结构化内容失败",
                file_path=file_path
            )
        
        return parse_result.metadata
    
    def extract_paragraphs(self, file_path: str) -> List[Dict[str, Any]]:
        """
        从 Word 文档中提取段落。
        
        参数:
            file_path: Word 文档路径
            
        返回:
            段落字典列表
            
        引发:
            ParsingError: 如果提取失败
        """
        structured_content = self.extract_structured_content(file_path)
        return structured_content.get("paragraphs", [])
    
    def get_document_properties(self, file_path: str) -> Dict[str, Any]:
        """
        从 Word 文档中提取文档属性。
        
        参数:
            file_path: Word 文档路径
            
        返回:
            包含文档属性的字典
        """
        properties = {}
        
        # 仅适用于 .docx 文件和 python-docx
        if (os.path.splitext(file_path)[1].lower() == '.docx' and 
            DocxDocument is not None):
            try:
                doc = DocxDocument(file_path)
                core_props = doc.core_properties
                
                properties = {
                    "title": core_props.title or "",
                    "author": core_props.author or "",
                    "subject": core_props.subject or "",
                    "keywords": core_props.keywords or "",
                    "created": core_props.created.isoformat() if core_props.created else "",
                    "modified": core_props.modified.isoformat() if core_props.modified else "",
                    "last_modified_by": core_props.last_modified_by or "",
                    "revision": str(core_props.revision) if core_props.revision else "",
                    "category": core_props.category or "",
                    "comments": core_props.comments or ""
                }
                
                # 移除空属性
                properties = {k: v for k, v in properties.items() if v}
                
            except Exception as e:
                logger.warning(f"从 {file_path} 提取文档属性失败: {str(e)}")
        
        return properties
    
    def combine_structured_content(self, structured_content: Dict[str, Any]) -> str:
        """
        将结构化的 Word 文档内容合并为单个文本字符串。
        
        参数:
            structured_content: 包含结构化内容的字典
            
        返回:
            合并的文本内容
        """
        if "full_text" in structured_content:
            return structured_content["full_text"]
        
        # 备用方法：合并段落
        text_parts = []
        if "paragraphs" in structured_content:
            for paragraph in structured_content["paragraphs"]:
                if isinstance(paragraph, dict) and "content" in paragraph:
                    text_parts.append(paragraph["content"])
                elif isinstance(paragraph, str):
                    text_parts.append(paragraph)
        
        return "\n\n".join(text_parts)