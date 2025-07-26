from argparse import ArgumentParser
import logging
import os
from typing import Optional, List
from datetime import datetime
from dotenv import load_dotenv
from mcp.server import Server
from mcp.server.fastmcp import FastMCP
from mcp.server.sse import SseServerTransport
from starlette.applications import Starlette
from starlette.requests import Request
from starlette.routing import Route, Mount
from starlette.responses import JSONResponse, Response  # 确保导入Response
import uvicorn
from pydantic import BaseModel
import fitz  # pymupdf
import markdown
from docx import Document as DocxDocument
import docx2txt  # 用于处理.doc和.docx文件
import olefile  # 用于处理OLE格式文件
import pypandoc  # 用于文档转换
import json
import faiss
import numpy as np
from sentence_transformers import SentenceTransformer
from langchain.text_splitter import RecursiveCharacterTextSplitter

# 配置日志
logging.basicConfig(
    level=logging.INFO, format="%(asctime)s - %(name)s - %(levelname)s  - %(message)s"
)
logger = logging.getLogger(__name__)

# 加载环境变量
load_dotenv()


class Config:
    """配置类"""

    HOST = os.getenv("HOST", "0.0.0.0")
    PORT = int(os.getenv("PORT", 8020))
    DEBUG = os.getenv("DEBUG", "False").lower() == "true"
    # 白名单目录配置，默认为当前目录下的docs文件夹
    DEFAULT_DOCS_DIR = os.path.join(os.path.dirname(__file__), "docs")
    ALLOWED_DIRS = os.getenv("ALLOWED_DIRS", DEFAULT_DOCS_DIR).split(",")
    EMBEDDING_MODEL = "paraphrase-multilingual-MiniLM-L12-v2"
    INDEX_DIR = os.path.join(os.path.dirname(__file__), "faiss_index")


# 初始化MCP - 添加协议版本
mcp = FastMCP("mcp_qa_server", protocol_version="1.0")

# 全局变量
embedding_model = None
faiss_index = None
document_store = {}
file_index_cache = {}  # 缓存文件索引信息，格式: {file_path: {"mtime": timestamp, "indexed": True}}

def init_embedding_model():
    """初始化嵌入模型"""
    global embedding_model
    if embedding_model is None:
        embedding_model = SentenceTransformer(Config.EMBEDDING_MODEL)
        logger.info(f"Embedding model {Config.EMBEDDING_MODEL} loaded")
    return embedding_model

def is_path_allowed(path: str) -> bool:
    """检查路径是否在白名单目录中"""
    abs_path = os.path.abspath(path)
    for allowed_dir in Config.ALLOWED_DIRS:
        if abs_path.startswith(os.path.abspath(allowed_dir)):
            return True
    return False

def is_file_indexed(file_path: str) -> bool:
    """
    检查文件是否已经建立索引，并且文件没有被修改
    
    Args:
        file_path: 文件路径
        
    Returns:
        bool: 如果文件已索引且未修改返回True
    """
    global file_index_cache
    
    if not os.path.exists(file_path):
        return False
        
    current_mtime = os.path.getmtime(file_path)
    cache_info = file_index_cache.get(file_path)
    
    if cache_info and cache_info.get("indexed"):
        cached_mtime = cache_info.get("mtime")
        # 如果文件修改时间没有变化，说明文件未被修改，索引仍然有效
        return current_mtime == cached_mtime
    
    return False

def auto_convert_doc_to_docx(file_path: str) -> dict:
    """
    自动将.doc文件转换为.docx格式
    
    Args:
        file_path: .doc文件路径
        
    Returns:
        dict: 转换结果，包含成功状态和临时文件路径
    """
    import tempfile
    import subprocess
    import shutil
    
    try:
        # 创建临时文件路径
        temp_dir = tempfile.mkdtemp()
        base_name = os.path.splitext(os.path.basename(file_path))[0]
        temp_docx_path = os.path.join(temp_dir, f"{base_name}_converted.docx")
        
        # 方法1: 尝试使用pypandoc转换
        try:
            logger.info(f"Attempting to convert {file_path} using pypandoc")
            pypandoc.convert_file(file_path, 'docx', outputfile=temp_docx_path)
            
            if os.path.exists(temp_docx_path) and os.path.getsize(temp_docx_path) > 0:
                return {
                    "success": True,
                    "converted_path": temp_docx_path,
                    "temp_dir": temp_dir,
                    "method": "pypandoc"
                }
        except Exception as e:
            logger.debug(f"pypandoc conversion failed: {e}")
        
        # 方法2: 尝试使用LibreOffice (如果安装了)
        libreoffice_paths = [
            '/Applications/LibreOffice.app/Contents/MacOS/soffice',  # macOS
            '/usr/bin/libreoffice',  # Linux
            '/usr/bin/soffice',  # Linux alternative
            'libreoffice',  # PATH
            'soffice'  # PATH alternative
        ]
        
        for soffice_path in libreoffice_paths:
            try:
                # 检查LibreOffice是否可用
                if not shutil.which(soffice_path) and not os.path.exists(soffice_path):
                    continue
                
                logger.info(f"Attempting to convert {file_path} using LibreOffice")
                
                # 使用LibreOffice转换
                cmd = [
                    soffice_path,
                    '--headless',
                    '--convert-to', 'docx',
                    '--outdir', temp_dir,
                    file_path
                ]
                
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
                
                # 检查转换后的文件
                expected_docx = os.path.join(temp_dir, f"{base_name}.docx")
                if os.path.exists(expected_docx) and os.path.getsize(expected_docx) > 0:
                    return {
                        "success": True,
                        "converted_path": expected_docx,
                        "temp_dir": temp_dir,
                        "method": "LibreOffice"
                    }
                    
            except (subprocess.TimeoutExpired, FileNotFoundError, Exception) as e:
                logger.debug(f"LibreOffice conversion failed with {soffice_path}: {e}")
                continue
        
        # 方法3: 尝试使用系统的textutil (macOS only)
        if os.name == 'posix' and subprocess.run(['which', 'textutil'], capture_output=True).returncode == 0:
            try:
                logger.info(f"Attempting to convert {file_path} using textutil (macOS)")
                temp_rtf = os.path.join(temp_dir, f"{base_name}.rtf")
                
                # 先转换为RTF
                subprocess.run([
                    'textutil', '-convert', 'rtf', 
                    '-output', temp_rtf,
                    file_path
                ], check=True, timeout=30)
                
                if os.path.exists(temp_rtf):
                    # 再用pypandoc将RTF转换为docx
                    pypandoc.convert_file(temp_rtf, 'docx', outputfile=temp_docx_path)
                    
                    if os.path.exists(temp_docx_path) and os.path.getsize(temp_docx_path) > 0:
                        return {
                            "success": True,
                            "converted_path": temp_docx_path,
                            "temp_dir": temp_dir,
                            "method": "textutil + pypandoc"
                        }
                        
            except Exception as e:
                logger.debug(f"textutil conversion failed: {e}")
        
        # 清理临时目录
        shutil.rmtree(temp_dir, ignore_errors=True)
        
        return {
            "success": False,
            "error": "All conversion methods failed",
            "tried_methods": ["pypandoc", "LibreOffice", "textutil"]
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": f"Conversion failed: {str(e)}"
        }

def parse_old_doc_file(file_path: str) -> dict:
    """
    尝试解析旧的.doc格式文件
    
    Args:
        file_path: .doc文件路径
        
    Returns:
        dict: 解析结果或错误信息
    """
    import re
    import unicodedata
    
    def clean_text(text: str) -> str:
        """清理和过滤文本内容"""
        if not text:
            return ""
        
        # 分行处理
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 过滤条件：
            # 1. 长度至少为3
            if len(line) < 3:
                continue
            
            # 2. 不包含太多非可打印字符
            printable_chars = sum(1 for c in line if c.isprintable())
            if printable_chars / len(line) < 0.7:
                continue
            
            # 3. 不是纯符号或数字
            if re.match(r'^[\W\d_]+$', line):
                continue
            
            # 4. 包含一定数量的字母或中文字符
            letters_or_chinese = len(re.findall(r'[a-zA-Z\u4e00-\u9fff]', line))
            if letters_or_chinese < 2:
                continue
            
            # 5. 不包含太多连续的特殊字符
            if re.search(r'[^\w\s\u4e00-\u9fff]{3,}', line):
                continue
            
            # 6. 不是常见的二进制标记
            binary_markers = ['\\x', '\x00', '\\0', '\r', '\\r', '\n', '\\n']
            if any(marker in line for marker in binary_markers):
                continue
            
            # 7. 过滤明显的乱码（大部分是非常用字符）
            unusual_chars = 0
            for char in line:
                if char.isprintable() and not char.isspace():
                    # 检查是否是常用字符
                    if not (char.isalnum() or 
                           char in '.,;:!?"\'()-[]{}/<>@#$%^&*+=|`~_\u4e00-\u9fff\uff00-\uffef'):
                        unusual_chars += 1
            
            if unusual_chars > len(line) * 0.3:  # 如果超过30%是不常用字符
                continue
            
            cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
    
    def is_meaningful_text(text: str) -> bool:
        """判断文本是否有意义"""
        if not text or len(text) < 50:
            return False
        
        # 检查是否包含足够的有意义字符
        meaningful_chars = len(re.findall(r'[a-zA-Z\u4e00-\u9fff]', text))
        if meaningful_chars < 20:
            return False
        
        # 检查是否包含常见的中文或英文单词
        common_patterns = [
            r'\b(?:the|and|or|of|to|in|for|with|by|from|as|at|on|be|have|do|will|would|could|should)\b',
            r'[一-鿿]{2,}',  # 中文词组
            r'\b[a-zA-Z]{3,}\b'  # 英文单词
        ]
        
        for pattern in common_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                return True
        
        return False
    
    def is_mostly_garbage(text: str) -> bool:
        """检测文本是否主要是乱码"""
        if not text or len(text) < 10:
            return True
        
        # 检查可读字符的比例
        printable_ratio = sum(1 for c in text if c.isprintable()) / len(text)
        if printable_ratio < 0.8:
            return True
        
        # 检查是否包含太多单字符"单词"
        words = text.split()
        single_char_words = sum(1 for word in words if len(word) == 1)
        if len(words) > 0 and single_char_words / len(words) > 0.7:
            return True
        
        # 检查是否包含太多不常见的字符组合
        unusual_patterns = 0
        for word in words[:50]:  # 只检查前50个单词
            if len(word) <= 6 and not re.match(r'^[a-zA-Z\u4e00-\u9fff\d\s,.!?;:"()\[\]{}/<>@#$%^&*+=|`~_-]+$', word):
                unusual_patterns += 1
        
        if len(words) > 0 and unusual_patterns / min(len(words), 50) > 0.5:
            return True
        
        return False
    
    try:
        # 方法1: 检查是否是OLE文件
        if olefile.isOleFile(file_path):
            logger.info(f"Detected OLE format .doc file: {file_path}")
            
            # 尝试使用strings命令提取文本
            try:
                import subprocess
                # 使用strings命令，只提取较长的字符串
                result = subprocess.run(
                    ['strings', '-n', '4', file_path],  # 至少4个字符
                    capture_output=True, 
                    text=True, 
                    timeout=30
                )
                if result.returncode == 0 and result.stdout.strip():
                    raw_text = result.stdout
                    cleaned_text = clean_text(raw_text)
                    
                    # 检查是否主要是乱码
                    if is_mostly_garbage(cleaned_text):
                        logger.info("Extracted text appears to be mostly garbage")
                    elif is_meaningful_text(cleaned_text):
                        return {
                            "success": True,
                            "text": cleaned_text,
                            "method": "strings command extraction"
                        }
            except (subprocess.TimeoutExpired, FileNotFoundError, Exception) as e:
                logger.debug(f"strings command failed: {e}")
        
        # 方法2: 直接读取并清理
        try:
            with open(file_path, 'rb') as f:
                # 读取二进制数据
                raw_data = f.read()
                
                # 尝试不同的编码
                for encoding in ['utf-8', 'gbk', 'gb2312', 'latin1', 'cp1252']:
                    try:
                        text = raw_data.decode(encoding, errors='ignore')
                        cleaned_text = clean_text(text)
                        
                        # 检查是否主要是乱码
                        if is_mostly_garbage(cleaned_text):
                            continue
                        elif is_meaningful_text(cleaned_text):
                            return {
                                "success": True,
                                "text": cleaned_text,
                                "method": f"binary read with {encoding} encoding"
                            }
                    except Exception:
                        continue
        except Exception as e:
            logger.debug(f"Binary read failed: {e}")
        
        # 方法3: 使用正则表达式提取可能的文本块
        try:
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                
            # 查找可能的文本区域
            text_chunks = []
            for encoding in ['utf-8', 'gbk', 'latin1']:
                try:
                    decoded = raw_data.decode(encoding, errors='ignore')
                    # 提取连续的可读字符
                    matches = re.findall(r'[\u4e00-\u9fff\w\s,.!?;:"()\[\]{}/<>@#$%^&*+=|`~_-]{10,}', decoded)
                    for match in matches:
                        cleaned = clean_text(match)
                        if cleaned and len(cleaned) > 20:
                            text_chunks.append(cleaned)
                except Exception:
                    continue
            
            if text_chunks:
                combined_text = '\n\n'.join(text_chunks)
                if is_meaningful_text(combined_text):
                    return {
                        "success": True,
                        "text": combined_text,
                        "method": "regex text extraction"
                    }
        except Exception as e:
            logger.debug(f"Regex extraction failed: {e}")
        
        return {
            "success": False,
            "error": "Unable to extract meaningful text from .doc file",
            "message": "旧的.doc文件格式无法解析，或文件内容主要为二进制数据",
            "suggestions": [
                "1. 使用Microsoft Word打开文件，然后另存为.docx格式",
                "2. 使用LibreOffice Writer打开文件，然后导出为.docx或.pdf格式",
                "3. 在线转换工具：使用在线DOC到DOCX转换器",
                "4. 如果文件内容不多，可以手动复制粘贴到新的.docx文件中"
            ],
            "technical_note": "旧.doc格式使用二进制OLE结构，需要专门的解析器才能正确解析"
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": f"Failed to parse .doc file: {str(e)}"
        }

def mark_file_indexed(file_path: str):
    """
    标记文件已建立索引
    
    Args:
        file_path: 文件路径
    """
    global file_index_cache
    
    if os.path.exists(file_path):
        file_index_cache[file_path] = {
            "mtime": os.path.getmtime(file_path),
            "indexed": True
        }
        logger.info(f"Marked file as indexed: {file_path}")

def get_file_from_index(file_path: str) -> dict:
    """
    从索引中获取文件信息
    
    Args:
        file_path: 文件路径
        
    Returns:
        dict: 文件在索引中的信息
    """
    global document_store
    
    file_chunks = []
    for i, doc_info in enumerate(document_store):
        if doc_info.get("source") == file_path:
            file_chunks.append({
                "chunk_id": doc_info.get("metadata", {}).get("chunk_id", i),
                "content": doc_info.get("content", ""),
                "metadata": doc_info.get("metadata", {})
            })
    
    return {
        "file_path": file_path,
        "chunks": [chunk["content"] for chunk in file_chunks],
        "chunk_count": len(file_chunks),
        "chunks_detail": file_chunks,
        "from_cache": True
    }



# 定义工具调用参数模型
class ReadFileParams(BaseModel):
    file_name: str  # 文件名或路径
    force_reindex: bool = False  # 是否强制重新分析和构建索引

class ListDirParams(BaseModel):
    directory: str

class GetMtimeParams(BaseModel):
    file_path: str

class ParsePdfParams(BaseModel):
    file_path: str

class ParseDocxParams(BaseModel):
    file_path: str

class ParseMdParams(BaseModel):
    file_path: str

class SearchDocumentsParams(BaseModel):
    query: str
    top_k: Optional[int] = 3


@mcp.tool()
async def list_dir(params: ListDirParams):
    """
    列出目录中的文件和子目录
    
    Args:
        directory: 要列出的目录路径
        
    Returns:
        目录内容列表或错误信息
    """
    try:
        if not is_path_allowed(params.directory):
            return {"error": "Access denied", "message": "Directory not in allowed list"}
            
        if not os.path.exists(params.directory):
            return {"error": "Directory not found", "path": params.directory}
            
        if not os.path.isdir(params.directory):
            return {"error": "Not a directory", "path": params.directory}
            
        items = []
        for item in os.listdir(params.directory):
            item_path = os.path.join(params.directory, item)
            is_dir = os.path.isdir(item_path)
            size = os.path.getsize(item_path) if not is_dir else 0
            mtime = os.path.getmtime(item_path)
            
            items.append({
                "name": item,
                "type": "directory" if is_dir else "file",
                "size": size,
                "modified_time": datetime.fromtimestamp(mtime).isoformat()
            })
            
        return {"directory": params.directory, "items": items}
        
    except PermissionError:
        return {"error": "Permission denied", "path": params.directory}
    except Exception as e:
        logger.exception(f"Error listing directory: {params.directory}")
        return {"error": "Unexpected error", "details": str(e)}

@mcp.tool()
async def get_mtime(params: GetMtimeParams):
    """
    获取文件的修改时间
    
    Args:
        file_path: 文件路径
        
    Returns:
        文件修改时间或错误信息
    """
    try:
        if not os.path.exists(params.file_path):
            return {"error": "File not found", "path": params.file_path}
            
        mtime = os.path.getmtime(params.file_path)
        return {
            "file_path": params.file_path,
            "modified_time": datetime.fromtimestamp(mtime).isoformat(),
            "timestamp": mtime
        }
        
    except Exception as e:
        logger.exception(f"Error getting mtime for: {params.file_path}")
        return {"error": "Unexpected error", "details": str(e)}

@mcp.tool()
async def get_current_time():
    """
    获取当前时间
    
    Returns:
        当前时间信息
    """
    now = datetime.now()
    return {
        "current_time": now.isoformat(),
        "timestamp": now.timestamp(),
        "formatted": now.strftime("%Y年%m月%d日 %H:%M:%S"),
        "date": now.strftime("%Y-%m-%d"),
        "time": now.strftime("%H:%M:%S")
    }
@mcp.tool()
async def parse_pdf(params: ParsePdfParams):
    """
    解析PDF文档内容
    
    Args:
        file_path: PDF文件路径
        
    Returns:
        解析后的文档内容和结构化信息
    """
    try:
        if not os.path.exists(params.file_path):
            return {"error": "File not found", "path": params.file_path}
            
        if not params.file_path.lower().endswith('.pdf'):
            return {"error": "Not a PDF file", "path": params.file_path}
            
        doc = fitz.open(params.file_path)
        pages_content = []
        full_text = ""
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            page_text = page.get_text()
            # 清理文本
            cleaned_text = page_text.replace('\x0c', '').strip()
            if cleaned_text:
                pages_content.append({
                    "page_number": page_num + 1,
                    "content": cleaned_text
                })
                full_text += cleaned_text + "\n\n"
        
        doc.close()
        
        # 段落分割
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separators=["\n\n", "\n", "。", "！", "？", ";", ":", ".", "!", "?"]
        )
        chunks = text_splitter.split_text(full_text)
        
        return {
            "file_path": params.file_path,
            "total_pages": len(pages_content),
            "full_text": full_text,
            "pages": pages_content,
            "chunks": chunks,
            "chunk_count": len(chunks)
        }
        
    except Exception as e:
        logger.exception(f"Error parsing PDF: {params.file_path}")
        return {"error": "PDF parsing failed", "details": str(e), "path": params.file_path}

@mcp.tool()
async def parse_docx(params: ParseDocxParams):
    """
    解析Word文档内容（支持.docx和.doc格式）
    
    Args:
        file_path: Word文档路径
        
    Returns:
        解析后的文档内容和结构化信息
    """
    try:
        if not os.path.exists(params.file_path):
            return {"error": "File not found", "path": params.file_path}
            
        file_ext = params.file_path.lower()
        if not file_ext.endswith(('.docx', '.doc')):
            return {"error": "Not a Word document", "path": params.file_path}
        
        # 使用docx2txt处理所有Word文档（支持.doc和.docx）
        try:
            # 对于.doc文件，首先尝试自动转换为.docx
            if file_ext.endswith('.doc') and not file_ext.endswith('.docx'):
                logger.info(f"Attempting to auto-convert .doc file: {params.file_path}")
                
                # 尝试自动转换
                conversion_result = auto_convert_doc_to_docx(params.file_path)
                
                if conversion_result.get("success"):
                    logger.info(f"Successfully converted using {conversion_result['method']}")
                    
                    try:
                        # 使用转换后的.docx文件进行解析
                        doc = DocxDocument(conversion_result["converted_path"])
                        paragraphs = []
                        full_text = ""
                        
                        for i, paragraph in enumerate(doc.paragraphs):
                            if paragraph.text.strip():
                                paragraphs.append({
                                    "paragraph_number": i + 1,
                                    "content": paragraph.text.strip()
                                })
                                full_text += paragraph.text.strip() + "\n\n"
                        
                        # 清理临时文件
                        import shutil
                        shutil.rmtree(conversion_result["temp_dir"], ignore_errors=True)
                        
                        if not full_text.strip():
                            return {
                                "error": "Empty converted document",
                                "message": "转换后的文档为空",
                                "path": params.file_path,
                                "conversion_method": conversion_result["method"]
                            }
                        
                        # 文本分块
                        text_splitter = RecursiveCharacterTextSplitter(
                            chunk_size=1000,
                            chunk_overlap=200,
                            separators=["\n\n", "\n", "。", "！", "？", ";", ":", ".", "!", "?"]
                        )
                        chunks = text_splitter.split_text(full_text)
                        
                        return {
                            "file_path": params.file_path,
                            "file_type": "doc",
                            "total_paragraphs": len(paragraphs),
                            "full_text": full_text,
                            "paragraphs": paragraphs,
                            "chunks": chunks,
                            "chunk_count": len(chunks),
                            "conversion_method": conversion_result["method"],
                            "note": "Successfully auto-converted from .doc to .docx format"
                        }
                        
                    except Exception as docx_parse_error:
                        # 清理临时文件
                        import shutil
                        shutil.rmtree(conversion_result["temp_dir"], ignore_errors=True)
                        
                        logger.error(f"Failed to parse converted docx: {docx_parse_error}")
                        # 如果转换成功但解析失败，回退到原来的方法
                        pass
                else:
                    logger.info(f"Auto-conversion failed: {conversion_result.get('error')}")
                    # 转换失败，继续使用原来的方法
                
                # 如果自动转换失败，尝试使用专门的.doc解析方法
                doc_result = parse_old_doc_file(params.file_path)
                if doc_result.get("success"):
                    full_text = doc_result["text"]
                    logger.info(f"Successfully parsed .doc file using {doc_result['method']}")
                else:
                    # 如果专门方法失败，尝试docx2txt
                    try:
                        full_text = docx2txt.process(params.file_path)
                        if not full_text.strip():
                            return {
                                "error": "Unable to parse DOC file",
                                "message": "旧的.doc文件格式无法解析，所有解析方法都失败了",
                                "path": params.file_path,
                                "tried_methods": ["OLE extraction", "text extraction", "encoding variants", "docx2txt"],
                                "suggestions": [
                                    "1. 使用Microsoft Word打开文件，然后另存为.docx格式",
                                    "2. 使用LibreOffice Writer打开文件，然后导出为.docx或.pdf格式",
                                    "3. 在线转换工具：使用在线DOC到DOCX转换器"
                                ],
                                "details": doc_result.get("error", "Unknown error")
                            }
                    except Exception as docx2txt_error:
                        return {
                            "error": "DOC parsing failed",
                            "message": f"无法解析.doc文件: {str(docx2txt_error)}",
                            "path": params.file_path,
                            "suggestion": "请检查文件是否损坏或尝试将其转换为.docx格式",
                            "details": str(docx2txt_error)
                        }
            else:
                # 对于.docx文件，直接使用docx2txt
                full_text = docx2txt.process(params.file_path)
            
            if not full_text.strip():
                return {
                    "error": "Empty document",
                    "message": "文档为空或无法提取文本内容",
                    "path": params.file_path
                }
            
            # 将文本按段落分割
            paragraphs = []
            lines = full_text.split('\n')
            paragraph_num = 1
            
            for line in lines:
                line = line.strip()
                if line:  # 非空行
                    paragraphs.append({
                        "paragraph_number": paragraph_num,
                        "content": line
                    })
                    paragraph_num += 1
            
            # 清理和格式化文本
            cleaned_text = '\n\n'.join([p['content'] for p in paragraphs])
            
            # 段落分割
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=200,
                separators=["\n\n", "\n", "。", "！", "？", ";", ":", ".", "!", "?"]
            )
            chunks = text_splitter.split_text(cleaned_text)
            
            file_type = "docx" if file_ext.endswith('.docx') else "doc"
            
            return {
                "file_path": params.file_path,
                "file_type": file_type,
                "total_paragraphs": len(paragraphs),
                "full_text": cleaned_text,
                "paragraphs": paragraphs,
                "chunks": chunks,
                "chunk_count": len(chunks)
            }
            
        except Exception as docx2txt_error:
            # 如果是.docx文件，尝试使用python-docx作为备选
            if file_ext.endswith('.docx'):
                try:
                    logger.info(f"docx2txt failed, trying python-docx for {params.file_path}")
                    doc = DocxDocument(params.file_path)
                    paragraphs = []
                    full_text = ""
                    
                    for i, paragraph in enumerate(doc.paragraphs):
                        if paragraph.text.strip():
                            paragraphs.append({
                                "paragraph_number": i + 1,
                                "content": paragraph.text.strip()
                            })
                            full_text += paragraph.text.strip() + "\n\n"
                    
                    # 段落分割
                    text_splitter = RecursiveCharacterTextSplitter(
                        chunk_size=1000,
                        chunk_overlap=200,
                        separators=["\n\n", "\n", "。", "！", "？", ";", ":", ".", "!", "?"]
                    )
                    chunks = text_splitter.split_text(full_text)
                    
                    return {
                        "file_path": params.file_path,
                        "file_type": "docx",
                        "total_paragraphs": len(paragraphs),
                        "full_text": full_text,
                        "paragraphs": paragraphs,
                        "chunks": chunks,
                        "chunk_count": len(chunks),
                        "note": "Parsed using python-docx as fallback"
                    }
                except Exception as python_docx_error:
                    return {
                        "error": "Both parsing methods failed",
                        "message": "docx2txt和python-docx都无法解析此文件",
                        "path": params.file_path,
                        "details": {
                            "docx2txt_error": str(docx2txt_error),
                            "python_docx_error": str(python_docx_error)
                        }
                    }
            else:
                # .doc文件只能使用docx2txt
                return {
                    "error": "DOC parsing failed",
                    "message": f"无法解析.doc文件: {str(docx2txt_error)}",
                    "path": params.file_path,
                    "suggestion": "请检查文件是否损坏或尝试将其转换为.docx格式",
                    "details": str(docx2txt_error)
                }
        
    except Exception as e:
        logger.exception(f"Error parsing Word document: {params.file_path}")
        return {"error": "Document parsing failed", "details": str(e), "path": params.file_path}

@mcp.tool()
async def parse_md(params: ParseMdParams):
    """
    解析Markdown文档内容
    
    Args:
        file_path: Markdown文件路径
        
    Returns:
        解析后的文档内容和结构化信息
    """
    try:
        if not os.path.exists(params.file_path):
            return {"error": "File not found", "path": params.file_path}
            
        if not params.file_path.lower().endswith(('.md', '.markdown')):
            return {"error": "Not a Markdown file", "path": params.file_path}
            
        with open(params.file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # 解析markdown
        md = markdown.Markdown(extensions=['meta', 'toc'])
        html_content = md.convert(md_content)
        
        # 段落分割
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separators=["\n\n", "\n", "。", "！", "？", ";", ":", ".", "!", "?"]
        )
        chunks = text_splitter.split_text(md_content)
        
        return {
            "file_path": params.file_path,
            "raw_content": md_content,
            "html_content": html_content,
            "metadata": getattr(md, 'Meta', {}),
            "toc": getattr(md, 'toc', ''),
            "chunks": chunks,
            "chunk_count": len(chunks)
        }
        
    except Exception as e:
        logger.exception(f"Error parsing Markdown: {params.file_path}")
        return {"error": "Markdown parsing failed", "details": str(e), "path": params.file_path}
@mcp.tool()
async def search_documents(params: SearchDocumentsParams):
    """
    在文档索引中搜索相关内容
    
    Args:
        query: 搜索查询
        top_k: 返回前K个结果，默认3
        
    Returns:
        搜索结果列表
    """
    try:
        global faiss_index, document_store, embedding_model
        
        # 初始化嵌入模型
        if embedding_model is None:
            init_embedding_model()
        
        # 加载或初始化FAISS索引
        if faiss_index is None:
            index_path = os.path.join(Config.INDEX_DIR, "index.faiss")
            store_path = os.path.join(Config.INDEX_DIR, "index.pkl")
            
            if os.path.exists(index_path) and os.path.exists(store_path):
                faiss_index = faiss.read_index(index_path)
                with open(store_path, 'rb') as f:
                    import pickle
                    document_store = pickle.load(f)
                logger.info("Loaded existing FAISS index")
            else:
                return {"error": "No index found", "message": "请先创建文档索引"}
        
        # 对查询进行向量化
        query_embedding = embedding_model.encode([params.query])
        query_vector = np.array(query_embedding).astype('float32')
        
        # 搜索
        distances, indices = faiss_index.search(query_vector, params.top_k)
        
        results = []
        for i, (distance, idx) in enumerate(zip(distances[0], indices[0])):
            if idx < len(document_store):
                doc_info = document_store[idx]
                results.append({
                    "rank": i + 1,
                    "score": float(1 / (1 + distance)),  # 转换为相似度分数
                    "distance": float(distance),
                    "content": doc_info.get("content", ""),
                    "source": doc_info.get("source", ""),
                    "metadata": doc_info.get("metadata", {})
                })
        
        return {
            "query": params.query,
            "total_results": len(results),
            "results": results
        }
        
    except Exception as e:
        logger.exception(f"Error searching documents: {params.query}")
        return {"error": "Search failed", "details": str(e)}

@mcp.tool()
async def build_document_index(params: dict):
    """
    构建文档向量索引
    
    Args:
        params: 可包含files参数（文件路径列表）或directory参数
        
    Returns:
        索引构建结果
    """
    try:
        global faiss_index, document_store, embedding_model
        
        # 初始化嵌入模型
        if embedding_model is None:
            init_embedding_model()
        
        documents = []
        
        # 如果指定了文件列表
        files = params.get("files", [])
        directory = params.get("directory", "")
        
        if directory and is_path_allowed(directory):
            # 扫描目录中的文档文件
            for root, _, files_in_dir in os.walk(directory):
                for file in files_in_dir:
                    if file.lower().endswith(('.pdf', '.docx', '.doc', '.md', '.txt')):
                        files.append(os.path.join(root, file))
        
        for file_path in files:
            if not is_path_allowed(file_path):
                continue
                
            if not os.path.exists(file_path):
                continue
                
            # 根据文件类型解析
            if file_path.lower().endswith('.pdf'):
                result = await parse_pdf(ParsePdfParams(file_path=file_path))
                if "chunks" in result:
                    for i, chunk in enumerate(result["chunks"]):
                        documents.append({
                            "content": chunk,
                            "source": file_path,
                            "metadata": {
                                "type": "pdf",
                                "chunk_id": i,
                                "total_chunks": result["chunk_count"]
                            }
                        })
            elif file_path.lower().endswith(('.docx', '.doc')):
                result = await parse_docx(ParseDocxParams(file_path=file_path))
                if "chunks" in result:
                    for i, chunk in enumerate(result["chunks"]):
                        documents.append({
                            "content": chunk,
                            "source": file_path,
                            "metadata": {
                                "type": "docx",
                                "chunk_id": i,
                                "total_chunks": result["chunk_count"]
                            }
                        })
            elif file_path.lower().endswith(('.md', '.markdown')):
                result = await parse_md(ParseMdParams(file_path=file_path))
                if "chunks" in result:
                    for i, chunk in enumerate(result["chunks"]):
                        documents.append({
                            "content": chunk,
                            "source": file_path,
                            "metadata": {
                                "type": "markdown",
                                "chunk_id": i,
                                "total_chunks": result["chunk_count"]
                            }
                        })
            elif file_path.lower().endswith('.txt'):
                # 处理文本文件
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                    
                text_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=1000,
                    chunk_overlap=200
                )
                chunks = text_splitter.split_text(content)
                
                for i, chunk in enumerate(chunks):
                    documents.append({
                        "content": chunk,
                        "source": file_path,
                        "metadata": {
                            "type": "txt",
                            "chunk_id": i,
                            "total_chunks": len(chunks)
                        }
                    })
        
        if not documents:
            return {"error": "No documents found", "message": "没有找到可处理的文档"}
        
        # 生成嵌入向量
        contents = [doc["content"] for doc in documents]
        embeddings = embedding_model.encode(contents, batch_size=32, show_progress_bar=True)
        
        # 创建FAISS索引
        dimension = embeddings.shape[1]
        faiss_index = faiss.IndexFlatL2(dimension)
        faiss_index.add(embeddings.astype('float32'))
        
        # 保存文档信息
        document_store = documents
        
        # 确保索引目录存在
        os.makedirs(Config.INDEX_DIR, exist_ok=True)
        
        # 保存索引
        index_path = os.path.join(Config.INDEX_DIR, "index.faiss")
        store_path = os.path.join(Config.INDEX_DIR, "index.pkl")
        
        faiss.write_index(faiss_index, index_path)
        with open(store_path, 'wb') as f:
            import pickle
            pickle.dump(document_store, f)
        
        logger.info(f"Built index with {len(documents)} documents")
        
        return {
            "message": "索引构建完成",
            "total_documents": len(documents),
            "index_dimension": dimension,
            "index_path": index_path,
            "files_processed": len(files)
        }
        
    except Exception as e:
        logger.exception("Error building document index")
        return {"error": "Index building failed", "details": str(e)}

@mcp.tool()
async def read_file(params: ReadFileParams):
    """
    读取并解析指定文件，支持多种格式并构建向量索引
    白名单路径下的文件不受限制，其他路径的文件需要在白名单目录中
    如果文件已经构建索引且未修改，则优先使用缓存的索引

    Args:
        file_name: 要读取的文件名或相对路径
        force_reindex: 是否强制重新分析和构建索引

    Returns:
        解析后的文件内容、结构化信息和向量索引结果
    """
    file_name = params.file_name
    force_reindex = params.force_reindex
    found_path = None
    
    # 查找文件
    if os.path.exists(file_name):
        # 文件存在，检查是否在白名单路径下
        if is_path_allowed(file_name):
            found_path = file_name
        else:
            return {
                "error": "Access denied",
                "message": f"文件 '{file_name}' 不在白名单目录中",
                "file_path": file_name
            }
    else:
        # 在白名单目录中查找文件
        for allowed_dir in Config.ALLOWED_DIRS:
            allowed_dir = allowed_dir.strip()
            
            potential_paths = [
                os.path.join(allowed_dir, file_name),
                os.path.join(allowed_dir, os.path.basename(file_name)),
                file_name if file_name.startswith(allowed_dir) else None,
            ]
            
            for potential_path in potential_paths:
                if potential_path and os.path.exists(potential_path) and os.path.isfile(potential_path):
                    found_path = potential_path
                    break
            
            if found_path:
                break
    
    if not found_path:
        # 如果没找到，列出白名单目录中的可用文件
        available_files = []
        for allowed_dir in Config.ALLOWED_DIRS:
            allowed_dir = allowed_dir.strip()
            if os.path.exists(allowed_dir):
                try:
                    for root, dirs, files in os.walk(allowed_dir):
                        for file in files:
                            rel_path = os.path.relpath(os.path.join(root, file), '.')
                            available_files.append(rel_path)
                except Exception:
                    pass
        
        return {
            "error": "File not found",
            "message": f"文件 '{file_name}' 在白名单目录中未找到",
            "searched_directories": Config.ALLOWED_DIRS,
            "searched_for": file_name,
            "available_files": sorted(available_files[:15])
        }
    
    try:
        logger.info(f"Reading and parsing file: {found_path}")
        
        # 获取文件扩展名
        file_ext = os.path.splitext(found_path)[1].lower()
        
        # 检查是否已有索引（除非强制重新分析）
        if not force_reindex and is_file_indexed(found_path):
            logger.info(f"Using cached index for file: {found_path}")
            cached_result = get_file_from_index(found_path)
            
            # 根据文件类型添加相应的类型信息
            if file_ext == '.pdf':
                cached_result["file_type"] = "pdf"
                cached_result["total_pages"] = "N/A (from cache)"
            elif file_ext in ['.docx', '.doc']:
                file_type = "docx" if file_ext == '.docx' else "doc"
                cached_result["file_type"] = file_type
                cached_result["total_paragraphs"] = "N/A (from cache)"
            elif file_ext in ['.md', '.markdown']:
                cached_result["file_type"] = "markdown"
                cached_result["raw_content"] = "N/A (from cache)"
                cached_result["html_content"] = "N/A (from cache)"
            else:
                cached_result["file_type"] = "text"
                cached_result["content"] = "N/A (from cache)"
            
            cached_result["vector_index"] = {"message": "使用现有索引，未重新构建"}
            return cached_result
        
        # 根据文件类型进行解析
        if file_ext == '.pdf':
            result = await parse_pdf(ParsePdfParams(file_path=found_path))
            if "error" not in result:
                # 构建向量索引
                index_result = await build_document_index({"files": [found_path]})
                result["vector_index"] = index_result
                # 标记文件已索引
                mark_file_indexed(found_path)
            return result
            
        elif file_ext in ['.docx', '.doc']:
            result = await parse_docx(ParseDocxParams(file_path=found_path))
            if "error" not in result:
                # 构建向量索引
                index_result = await build_document_index({"files": [found_path]})
                result["vector_index"] = index_result
                # 标记文件已索引
                mark_file_indexed(found_path)
            return result
            
        elif file_ext in ['.md', '.markdown']:
            result = await parse_md(ParseMdParams(file_path=found_path))
            if "error" not in result:
                # 构建向量索引
                index_result = await build_document_index({"files": [found_path]})
                result["vector_index"] = index_result
                # 标记文件已索引
                mark_file_indexed(found_path)
            return result
            
        elif file_ext in ['.txt', '.py', '.js', '.json', '.yaml', '.yml', '.xml', '.csv']:
            # 文本文件处理
            with open(found_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # 文本分块
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=200,
                separators=["\n\n", "\n", "。", "！", "？", ";", ":", ".", "!", "?"]
            )
            chunks = text_splitter.split_text(content)
            
            # 构建向量索引
            index_result = await build_document_index({"files": [found_path]})
            
            # 标记文件已索引
            mark_file_indexed(found_path)
            
            return {
                "file_path": found_path,
                "file_type": "text",
                "content": content,
                "chunks": chunks,
                "chunk_count": len(chunks),
                "encoding": "utf-8",
                "vector_index": index_result
            }
            
        else:
            # 不支持的文件类型
            return {
                "error": "Unsupported file type",
                "message": f"不支持的文件类型: {file_ext}",
                "file_path": found_path,
                "supported_types": [".pdf", ".docx", ".doc", ".md", ".markdown", ".txt", ".py", ".js", ".json", ".yaml", ".yml", ".xml", ".csv"]
            }
            
    except UnicodeDecodeError:
        return {
            "error": "Encoding error",
            "message": "文件编码错误，无法使用UTF-8解码",
            "file_path": found_path
        }
    except Exception as e:
        logger.exception(f"Error reading file: {found_path}")
        return {"error": "Unexpected error", "details": str(e), "file_path": found_path}


def create_starlette_app(mcp_server: Server) -> Starlette:
    """
    创建基于Starlette的MCP服务器应用

    Args:
        mcp_server: MCP服务器实例

    Returns:
        Starlette应用实例
    """
    sse = SseServerTransport("/messages/")

    async def handle_sse(request: Request) -> Response:
        """处理SSE连接"""
        try:
            # 添加请求日志以帮助调试
            logger.debug(f"Received SSE connection from {request.client.host}")

            async with sse.connect_sse(
                request.scope,
                request.receive,
                request._send,
            ) as (read_stream, write_stream):
                # 修复：直接使用InitializationOptions对象，不进行解包
                await mcp_server.run(
                    read_stream,
                    write_stream,
                    mcp_server.create_initialization_options(),
                )

            # 成功完成返回200响应
            return Response(status_code=200)
        except Exception as e:
            logger.error(f"Error in SSE handler: {str(e)}")
            return JSONResponse({"error": "Internal server error"}, status_code=500)

    # 添加健康检查端点
    async def health_check(request: Request):
        return JSONResponse({"status": "ok", "service": "mcp-server"})

    return Starlette(
        debug=Config.DEBUG,
        routes=[
            Route("/sse", endpoint=handle_sse),
            Route("/health", endpoint=health_check),
            Mount("/messages/", app=sse.handle_post_message),
        ],
        on_startup=[lambda: logger.info("Server starting...")],
        on_shutdown=[lambda: logger.info("Server shutting down...")],
    )


def parse_arguments():
    """解析命令行参数"""
    parser = ArgumentParser(description="Run MCP SSE-based server")
    parser.add_argument("--host", default=Config.HOST, help="Host to bind to")
    parser.add_argument(
        "--port", type=int, default=Config.PORT, help="Port to listen on"
    )
    parser.add_argument("--debug", action="store_true", help="Enable debug mode")
    parser.add_argument("--allowed-dirs", help="Comma-separated list of allowed directories")
    return parser.parse_args()


if __name__ == "__main__":
    args = parse_arguments()
    # 更新配置
    Config.HOST = args.host
    Config.PORT = args.port
    Config.DEBUG = args.debug
    
    # 如果命令行指定了白名单目录，则更新配置
    if args.allowed_dirs:
        Config.ALLOWED_DIRS = args.allowed_dirs.split(",")

    # 确保白名单目录存在
    for allowed_dir in Config.ALLOWED_DIRS:
        if not os.path.exists(allowed_dir.strip()):
            logger.warning(f"Allowed directory does not exist: {allowed_dir.strip()}")
        else:
            logger.info(f"Allowed directory: {allowed_dir.strip()}")

    # 启动服务器
    mcp_server = mcp._mcp_server
    starlette_app = create_starlette_app(mcp_server)

    logger.info(f"Starting server on {Config.HOST}:{Config.PORT}")
    logger.info(f"Allowed directories: {Config.ALLOWED_DIRS}")

    uvicorn.run(
        starlette_app,
        host=Config.HOST,
        port=Config.PORT,
        log_level="info" if not Config.DEBUG else "debug",
    )
